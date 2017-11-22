# -*- coding=utf-8 -*-
from . import admin
from flask import render_template, flash, redirect, url_for, request, jsonify, send_from_directory, abort
from flask_login import login_required, current_user, login_user, logout_user
from forms import DownloadForm,LoginForm, RegistrationForm, EditRecordForm, ModifyForm, EditErecordForm
from ..models import User,Students
from .. import db,document
from xlwt import *
import os
import pdfkit
from docx.shared import Inches

@admin.route('/')
@login_required
def index():
    return render_template('admin/index.html')


@admin.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user is not None and user.verify_password(form.password.data):
            login_user(user)
            return redirect(url_for('admin.record'))
        flash(u'用户密码不正确')
    return render_template('admin/login.html', form=form)

@admin.route('/record', methods=['GET', 'POST'])
@login_required
def record():
    alist = Students.query.all()
    return render_template('admin/check.html', list=alist)

@admin.route('/modify/<int:id>', methods=['GET', 'POST'])
@login_required
def look(id):
    form = DownloadForm()
    if form.validate_on_submit():
        cord = Students.query.get_or_404(id)
        try:
            db.session.delete(cord)
            db.session.commit()
            return redirect(url_for('admin.record'))
        except:
            return redirect(url_for('admin.look', id=id))
    re = db.session.query(Students).filter(Students.id == id).one()
    docname = str(re.studentid) + '.docx'
    print docname
    document.add_heading(re.name, 0)
    '''
    document.add_picture('/var/www/registerflask/app/static/uploads/' + str(re.filename),
                            width=Inches(1.25))
    '''
    table = document.add_table(rows=2, cols=9, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u"性别"
    hdr_cells[1].text = u"学号"
    hdr_cells[2].text = u"手机号"
    hdr_cells[3].text = u"年级"
    hdr_cells[4].text = u"专业"
    hdr_cells[5].text = u"邮箱"
    hdr_cells[6].text = u"第一志愿"
    hdr_cells[7].text = u"第二志愿"
    hdr_cells[8].text = u"是否服从调剂"

    hdr_cells = table.rows[1].cells
    if str(re.sex) == '1':
        hdr_cells[0].text = u'男'
    else:
        hdr_cells[0].text = u'女'
    hdr_cells[1].text = re.studentid
    hdr_cells[2].text = re.phone
    if str(re.grade) == '1':
        hdr_cells[3].text = u'大一'
    elif str(re.grade) == '2':
        hdr_cells[3].text = u'大二'
    elif str(re.grade) == '3':
        hdr_cells[3].text = u'大三'
    elif str(re.grade) == '4':
        hdr_cells[3].text = u'大四'
    hdr_cells[4].text = re.major
    hdr_cells[5].text = re.email
    if str(re.question3) == '1':
        hdr_cells[6].text = u'电脑部'
    elif str(re.question3) == '2':
        hdr_cells[6].text = u'电器部'
    elif str(re.question3) == '3':
        hdr_cells[6].text = u'财外部'
    elif str(re.question3) == '4':
        hdr_cells[6].text = u'人资部'
    elif str(re.question3) == '5':
        hdr_cells[6].text = u'文宣部'

    document.add_heading(u' 请简述你选择该部门作为第一志愿的原因。', 1)
    document.add_paragraph(re.whywants1, style='IntenseQuote')

    if str(re.question4) == '1':
        hdr_cells[7].text = u'电脑部'
    elif str(re.question4) == '2':
        hdr_cells[7].text = u'电器部'
    elif str(re.question4) == '3':
        hdr_cells[7].text = u'财外部'
    elif str(re.question4) == '4':
        hdr_cells[7].text = u'人资部'
    elif str(re.question4) == '5':
        hdr_cells[7].text = u'文宣部'

    if re.question5 == 1:
        hdr_cells[8].text = u'服从'
    else:
        hdr_cells[8].text = u'不服从'
    document.add_heading(u' 请简述你选择该部门作为第二志愿的原因。', 1)
    document.add_paragraph(re.whywants2, style='IntenseQuote')

    document.add_heading(u'个人简介', 1)
    document.add_paragraph(re.intro, style='IntenseQuote')

    document.add_heading(u'请简述你想加入E志者协会的理由。', 1)
    document.add_paragraph(re.why, style='IntenseQuote')

    document.add_heading(u'请简述你能给E志者带来什么。', 1)
    document.add_paragraph(re.what, style='IntenseQuote')

    document.add_page_break()
    print docname
    document.save(docname)
    if os.path.isfile('app/static/doc'+docname):
        os.system("rm " + docname + "  app/static/doc/")
    os.system("mv " + docname + "  app/static/doc/")
    doc_url = url_for('static', filename=('doc/' + str(docname)))
    return render_template("admin/look.html", re = re,form = form,doc_url = doc_url)



@admin.route('/register', methods=['GET', 'POST'])
def register():
    register_key = 'zhucema'
    form = RegistrationForm()
    if form.validate_on_submit():
        if form.registerkey.data != register_key:
            flash(u'注册码不符，请返回重试')
            return redirect(url_for('admin.register'))
        else:
            if form.password.data != form.password2.data:
                flash(u'两次输入密码不一')
                return redirect(url_for('admin.register'))
            else:
                try:
                    user = User(username=form.username.data, password=form.password.data)
                    db.session.add(user)
                    db.session.commit()
                    return redirect(url_for('admin.login'))
                except:
                    db.session.rollback()
                    flash(u'用户名已存在')
    return render_template("admin/register.html", form=form)


@admin.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('admin.index'))


@admin.route('/download', methods=['GET', 'POST'])
@login_required
def download():
    re = db.session.query(Students)
    w = Workbook(encoding='utf-8')
    ws = w.add_sheet('baomingrecord')
    ws.write(0, 0, "序号")
    ws.write(0, 1, "姓名")
    ws.write(0, 2, "性别")
    ws.write(0, 3, "学号")
    ws.write(0, 4, "手机号码")
    ws.write(0, 5, "年级")
    ws.write(0, 6, "行政班")
    ws.write(0, 7, "邮箱")
    ws.write(0, 8, "之前有没有自己拆过机或是在他人的协助下拆过机")
    ws.write(0, 9, "如果我们给你拆的是一台款式老旧、结构古怪的机子你是否有信心和伙伴一起完成？")
    ws.write(0, 10, "当小e叫你拆如下图所示的这样一台机子时，你会从哪里入手开始拆？")
    ws.write(0, 11, "第一个被我们拆下的是哪部分重要硬件？")
    ws.write(0, 12, "笔记本电脑的键盘下的控制主板是用什么和显示屏相连的？？.25晚")
    ws.write(0, 13, "下列各图中所示的部件是显卡的是？")
    ws.write(0, 14, "11.24   13:00-15:00")
    ws.write(0, 15, "11.24   15:15-17:15")
    ws.write(0, 16, "11.24   18:30-20:30")
    ws.write(0, 17, "11.25   9:00-11:00")
    ws.write(0, 18, "11.25   13:00-15:00")
    ws.write(0, 19, "11.25   15:15-17:15")
    ws.write(0, 20, "11.25   18:30-20:30")
    ws.write(0, 21, "11.26   9:00-11:00")
    ws.write(0, 22, "11.26   13:00-15:00")
    ws.write(0, 23, "11.26   15:15-17:15")
    ws.write(0, 24, "备注")


    x = 1
    for stu in re:
        ws.write(x, 0, stu.id)
        ws.write(x, 1, stu.name)
        ws.write(x, 2, stu.sex)
        ws.write(x, 3, stu.studentid)
        ws.write(x, 4, stu.phone)
        ws.write(x, 5, stu.grade)
        ws.write(x, 6, stu.major)
        ws.write(x, 7, stu.email)
        ws.write(x, 8, stu.question1)
        ws.write(x, 9, stu.question2)
        ws.write(x, 10, stu.question3)
        ws.write(x, 11, stu.question4)
        ws.write(x, 12, stu.question5)
        ws.write(x, 13, stu.question6)
        ws.write(x, 14, stu.day242)
        ws.write(x, 15, stu.day243)
        ws.write(x, 16, stu.day244)
        ws.write(x, 17, stu.day251)
        ws.write(x, 18, stu.day252)
        ws.write(x, 19, stu.day253)
        ws.write(x, 20, stu.day254)
        ws.write(x, 21, stu.day261)
        ws.write(x, 22, stu.day262)
        ws.write(x, 23, stu.day263)
        ws.write(x, 24, stu.remarks)
        x += 1
    w.save('record.xls')
    # os.system("rm app/static/admin/record.xls")
    os.system("mv record.xls app/static/record/")
    # return send_from_directory("upload","/app/admin/static/record.xls", as_attachment=True)
    return render_template("admin/download.html")



